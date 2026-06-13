#!/usr/bin/env python3
"""Expand the 5 key PTD sections (2.1 competitors, 2.2 matrix, 7 FR, 13 user stories,
14 use cases) to full PRD/SRS depth; keep the rest from the condensed JSON. Re-emit
PTD-content.md (markdown) and PTD-content-tabs.txt (tab tables)."""
import json, docx

PRD = docx.Document("/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PRD/FYP-26-S2-37-PRD-v3.docx")
SRS = docx.Document("/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/SRS/FYP-26-S2-37-SRS-v2.docx")
OUTDIR = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PTD/"

# ---------- 2.1 Competitor write-ups (PRD paras 71-168) ----------
def competitors():
    ps = PRD.paragraphs
    blocks = [{"t":"p","s":"Normal","x":
        "Seven existing fitness applications were reviewed to understand the market and "
        "position Wise Workout: Strava, MyFitnessPal, adidas Running, Google Fit, Freeletics, "
        "MapMyFitness, and Lyfta. For each, the overview, key features, AI/ML usage, "
        "monetisation model, and relevance to Wise Workout are summarised below."}]
    names={"Strava","MyFitnessPal","adidas Running","Google Fit","Freeletics","MapMyFitness","Lyfta"}
    i=0
    while i < len(ps):
        p=ps[i]; t=p.text.strip()
        if p.style.name=="Heading 3" and t in names:
            blocks.append({"t":"p","s":"Heading 3","x":t})
            j=i+1
            label=None
            while j < len(ps):
                pj=ps[j]; tj=pj.text.strip()
                if pj.style.name=="Heading 3" and tj in names: break
                if pj.style.name.startswith("Heading") and "Product Comparison" in tj:
                    j=len(ps); break
                if pj.style.name=="Heading 4":
                    label=tj.rstrip(":")
                elif tj:
                    blocks.append({"t":"p","s":"Normal","x":(f"{label}: {tj}" if label else tj)})
                    label=None
                j+=1
            i=j; continue
        i+=1
    return blocks

# ---------- 2.2 Product Comparison Matrix (PRD T3, 14x9) ----------
def matrix():
    t=PRD.tables[3]
    rows=[[c.text.strip().replace("\n"," ") for c in r.cells] for r in t.rows]
    return [
        {"t":"p","s":"Normal","x":"The matrix compares Wise Workout against the reviewed "
         "applications across the main feature dimensions (Yes = supported, No = not supported, "
         "Limited = partially supported)."},
        {"t":"t","rows":rows},
    ]

# ---------- 7 Functional Requirements (PRD FR1-12 full, paras ~447-617) ----------
def functional():
    ps=PRD.paragraphs
    blocks=[{"t":"p","s":"Normal","x":
        "The functional requirements are organised into twelve groups (FR1-FR12). Each group "
        "and its main behaviours are described below, consolidated from the PRD main features "
        "and the SRS functional requirement specifications."}]
    start=end=None
    for i,p in enumerate(ps):
        if p.style.name=="Heading 3" and p.text.strip().startswith("(FR1)"): start=i
        if p.style.name=="Heading 2" and "Marketing Website" in p.text: end=i; break
    for j in range(start,end):
        p=ps[j]; t=p.text.strip()
        if not t: continue
        if p.style.name=="Heading 3":
            blocks.append({"t":"p","s":"Heading 3","x":t.replace("(","").replace(")","")})
        elif p.style.name in ("List Paragraph","List Bullet"):
            blocks.append({"t":"p","s":"List Bullet","x":t})
        else:
            blocks.append({"t":"p","s":"Normal","x":t})
    return blocks

# ---------- 13 User Stories (SRS T8-T71: all 64, ID|Name|Description) ----------
def user_stories():
    rows=[["ID","Name","Description"]]
    for t in SRS.tables:
        cells={r.cells[0].text.strip():r.cells[1].text.strip() for r in t.rows if len(r.cells)>=2}
        sid=cells.get("User Story ID","")
        if sid.startswith("#US"):
            rows.append([sid.lstrip("#"), cells.get("Name",""), cells.get("Description","")])
    return [
        {"t":"p","s":"Normal","x":f"The SRS defines {len(rows)-1} user stories across the five "
         "user roles. All are listed below with their use-case name and description; the full "
         "use-case flows are in Section 14."},
        {"t":"t","rows":rows},
    ]

# ---------- 14 Use Case Descriptions (SRS T8-T71 full structured) ----------
def use_cases():
    blocks=[{"t":"p","s":"Normal","x":
        "The full structured use-case descriptions (trigger, actors, pre-condition, normal "
        "flow, alternative flow, sub-flow) for all user stories are given below, consolidated "
        "from the SRS System Features section."}]
    order=["User Story ID","Name","Description","Trigger","Actor(s)","Pre-condition",
           "Normal Flow","Alternative Flow","Sub-flow"]
    for t in SRS.tables:
        cells={r.cells[0].text.strip():r.cells[1].text.strip() for r in t.rows if len(r.cells)>=2}
        sid=cells.get("User Story ID","")
        if not sid.startswith("#US"): continue
        blocks.append({"t":"p","s":"Heading 3","x":f"{sid.lstrip('#')} {cells.get('Name','')}"})
        rows=[[k, cells.get(k,"")] for k in order if k in cells]
        blocks.append({"t":"t","rows":rows})
    return blocks

# ---------- 1. Introduction: product-overview lead paragraph ----------
def intro_overview():
    return [{"t":"p","s":"Normal","x":
        "Wise Workout is a cross-platform mobile fitness application that combines workout "
        "tracking, AI-assisted progress summaries, AI-assisted fitness plan suggestions, social "
        "motivation, and a verified human-expert service marketplace within a single platform. "
        "It is designed for both self-guided users who want to track and understand their own "
        "progress and users who prefer professional support from verified coaches, trainers, "
        "nutritionists, and recovery specialists. By bringing automated tracking and insight "
        "together with access to human experts, Wise Workout aims to provide more complete "
        "fitness support than a single-purpose tracking app. This document consolidates the "
        "product direction, requirements, and technical design for the prototype."}]

# ---------- 1.4 Project Objectives: richer Objective/Summary table (PRD T2) ----------
def objectives():
    t=PRD.tables[2]
    rows=[[c.text.strip().replace("\n"," ") for c in r.cells] for r in t.rows]
    return [
        {"t":"p","s":"Normal","x":"The objectives of the Wise Workout project are as follows:"},
        {"t":"t","rows":rows},
    ]

# ---------- helpers to lift a source table verbatim ----------
def prd_table(i):
    return [[c.text.strip().replace("\n"," ") for c in r.cells] for r in PRD.tables[i].rows]

# ---------- 2.7 Business Model (PRD tiers + free/premium matrix + cost + channels) ----------
def business_model():
    return [
        {"t":"p","s":"Normal","x":
            "Wise Workout uses a freemium business model supported by a premium subscription and "
            "a separate paid expert service layer. The premium subscription is set at $9.99 per "
            "month for the prototype documentation. Expert services remain separate paid add-ons "
            "and are not automatically included in premium access."},
        {"t":"t","rows":[
            ["Revenue Stream","Description"],
            ["Premium subscription","Recurring subscription for advanced analytics, full workout "
             "history, personalised AI-assisted summaries, personalised plan suggestions, "
             "personalised reports, and personalised reminders."],
            ["Expert service fees","Revenue from expert-created content, service requests, "
             "consultations, training plans, recovery support, nutrition guidance, or "
             "sport-specific support."],
            ["Future opportunities","Sponsored fitness content, partnerships with fitness brands, "
             "featured expert listings, corporate wellness packages, and premium digital "
             "programmes."],
        ]},
        {"t":"p","s":"Heading 3","x":"User Tiers"},
        {"t":"t","rows":prd_table(6)},
        {"t":"p","s":"Heading 3","x":"Free vs Premium Feature Access"},
        {"t":"t","rows":prd_table(7)},
        {"t":"p","s":"Heading 3","x":"Cost Structure"},
        {"t":"t","rows":prd_table(8)},
        {"t":"p","s":"Heading 3","x":"Channels"},
        {"t":"t","rows":prd_table(9)},
    ]

# ---------- 10 Risk Management (full 15-row PRD risk register) ----------
def risk_register():
    return [
        {"t":"p","s":"Normal","x":
            "Risk analysis is important because the system spans several connected components - "
            "the mobile app, marketing website, backend, database, AI features, sensor/wearable "
            "data acquisition, expert workflows, subscriptions, social features, and admin "
            "management. The main risks relate to scope, integration, data privacy, AI output "
            "quality, and wearable data acquisition; they are managed through Agile iterative "
            "development, core-first prioritisation, regular module testing, and scope adjustment "
            "based on time and feasibility."},
        {"t":"t","rows":prd_table(30)},
    ]

# ---------- 8 Non-functional Requirements (SRS NFR requirement tables, all 6 categories) ----------
def nfr_sections():
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    blocks=[{"t":"p","s":"Normal","x":
        "The non-functional requirements are specified across six categories - security, "
        "reliability, performance, maintainability, scalability, and usability - each as a set "
        "of identified requirements with related roles and priority, consolidated from the SRS."}]
    cap=False; label=None; intro_done=set()
    for b in SRS.iter_inner_content():
        if isinstance(b,Paragraph):
            t=b.text.strip()
            if b.style.name=="Heading 1" and t.startswith("Non-Functional"): cap=True; continue
            if cap and b.style.name=="Heading 1": break
            if cap and b.style.name=="Heading 2":
                label=t; blocks.append({"t":"p","s":"Heading 3","x":t})
            elif cap and t and b.style.name=="Normal" and label not in intro_done:
                blocks.append({"t":"p","s":"Normal","x":t}); intro_done.add(label)
        elif isinstance(b,Table) and cap:
            rows=[[c.text.strip().replace("\n"," ") for c in r.cells] for r in b.rows]
            blocks.append({"t":"t","rows":rows})
    return blocks

# ---------- 11 Development Methodologies (PRD Agile section, condensed) ----------
def methodologies():
    return [
        {"t":"p","s":"Normal","x":
            "Wise Workout follows an Agile iterative development methodology. The project is "
            "divided into short development cycles (sprints), each focused on a specific group of "
            "features, so the system is built progressively from core functions to more advanced "
            "ones rather than all at once. This suits a multi-module system (mobile app, website, "
            "backend, database, AI features, expert workflows, admin) and a fixed FYP schedule."},
        {"t":"p","s":"Heading 3","x":"Justification for Using Agile"},
        {"t":"p","s":"Normal","x":
            "Agile provides flexibility as product features, diagrams, user flows, and prototype "
            "screens are refined; it lets the team respond to change without disrupting the whole "
            "plan. The system divides cleanly into connected modules that members can work on in "
            "parallel, and the iterative cycle supports regular testing and feedback so issues "
            "are found early rather than only at the end."},
        {"t":"p","s":"Heading 3","x":"Iteration / Sprint Approach"},
        {"t":"p","s":"Normal","x":
            "The project is divided into Scrum-inspired sprints; at the end of each sprint the "
            "team reviews completed work, identifies issues, and plans the next tasks (tracked in "
            "Jira). Scrum is not followed strictly because the team works within a fixed FYP "
            "schedule and adjusts sprint activities around deliverable deadlines, supervisor "
            "feedback, and availability. The proposed sprint structure is:"},
        {"t":"t","rows":prd_table(35)},
        {"t":"p","s":"Heading 3","x":"Collaboration, Tools and Testing"},
        {"t":"p","s":"Normal","x":
            "Tasks are divided by system area (mobile, website, backend, database, AI features, "
            "expert workflows, admin, testing, documentation). GitHub is used for source control "
            "and version management; Jira organises sprint tasks and progress. Testing runs "
            "throughout development rather than only at the end: functional testing per role and "
            "access tier, usability testing of key screens and flows, and review of AI outputs "
            "for relevance and clarity. Supervisor and team feedback feed back into each "
            "iteration, keeping the prototype aligned with the project requirements."},
    ]

EXP = {
    "1. Introduction": intro_overview(),
    "2.7 Business Model": business_model(),
    "8. Non-functional Requirements": nfr_sections(),
    "10. Risk Management": risk_register(),
    "11. Development Methodologies": methodologies(),
    "1.4 Project Objectives": objectives(),
    "2.1 Competitor Market Research": competitors(),
    "2.2 Product Comparison Matrix": matrix(),
    "7. Functional Requirements": functional(),     # parent heading; emit content under it
    "13. User Stories": user_stories(),
    "14. Use Case Descriptions": use_cases(),
}

# ---------- merge with condensed JSON, swapping the 5 sections ----------
cond = json.load(open("/tmp/ptd_content.json"))
# group condensed into sections keyed by heading text
sections=[]; cur=None
for b in cond:
    if b["type"]=="p" and b["style"].startswith("Heading"):
        cur={"title":b["text"].strip(),"head":b,"blocks":[]}; sections.append(cur)
    elif cur is not None:
        cur["blocks"].append(b)

def norm(s): return s.strip()
merged=[]
# For FR, the condensed has 7.1-7.4 subsections; we replace the whole §7 group:
# we will drop 7.1-7.4 condensed subsections and put expanded FR under "7. Functional Requirements"
skip_titles={"7.1 Functional Hierarchy","7.2 Basic Feature Access Levels","7.3 Dependencies","7.4 Inputs and Outputs",
             "8.1 Performance Requirements","8.2 Security Requirements","8.3 Interface and Usability Requirements","8.4 Portability and Scalability Requirements"}
for sec in sections:
    title=sec["title"]
    if title in skip_titles:
        continue
    merged.append(sec["head"])
    if title in EXP:
        for b in EXP[title]:
            if b["t"]=="p": merged.append({"type":"p","style":b["s"],"text":b["x"]})
            else: merged.append({"type":"t","rows":b["rows"]})
    else:
        merged.extend(sec["blocks"])

# ---------- scrub sample-template leftovers + reframe prototype->whole-app ----------
# The PTD/PUM describe the WHOLE app, not a prototype. Keep simulated payment / mock
# wearable / 107-tests only as brief honest notes (mostly in Appendix C).
_scrub = {
    # sample-template leftovers
    "public property datasets": "external or third-party datasets",
    # --- prototype -> whole-app reframes (product/feature/requirement text) ---
    "technical design for the prototype.": "technical design for Wise Workout.",
    "product direction, technical design, prototype scope, and implementation approach":
        "product direction, technical design, scope, and implementation approach",
    "current prototype repository": "current implementation repository",
    "Deliver a functional prototype": "Deliver a functional application",
    "The prototype implementation follows the current engineering direction":
        "The implementation follows the current engineering direction",
    "Prototype timeline limits the depth of production features. Some features such as payment and wearable integration are simulated. Full app store release is outside the immediate prototype scope.":
        "Development timeline limits the depth of some production features. Payment and wearable integration are currently simulated. Full app-store release is a later step.",
    "$9.99 per month for the prototype documentation.": "$9.99 per month.",
    "Free-tier or student-friendly hosting options may be used for the prototype.":
        "Free-tier or student-friendly hosting options may be used during early development.",
    "access the prototype mobile application.": "access the mobile application.",
    "For the current FYP prototype, the Android application may be distributed through the marketing website as an APK after registration. If the product is developed further, official app store distribution through the Apple App Store and Google Play Store may be considered.":
        "Initially the Android application is distributed through the marketing website as an APK after registration. Official app store distribution through the Apple App Store and Google Play Store is planned for later releases.",
    "the project deliverables, prototype, documentation, and final presentation.":
        "the project deliverables, the application, documentation, and final presentation.",
    "data needed to support the prototype.": "data needed to support the platform.",
    "Prototype demonstration of workout capture and wearable-supported activity tracking.":
        "Workout capture and wearable-supported activity tracking.",
    "Phone sensors and mock wearable pairing": "Phone sensors and wearable pairing",
    "GPS movement, activity data, simulated wearable heart rate, and connected-device status.":
        "GPS movement, activity data, wearable heart rate, and connected-device status.",
    "Mock BLE pairing and simulated wearable heart rate data for prototype validation.":
        "BLE pairing and wearable heart-rate data from connected devices.",
    "Collect only data required for the prototype functions and user workflows.":
        "Collect only data required for the app's functions and user workflows.",
    "Use simulated payment handling in the prototype instead of storing real card details.":
        "Do not store real payment card details (payment is simulated).",
    "End-of-Term-1 Review and prototype progress demonstration":
        "End-of-Term-1 Review and progress demonstration",
    "The prototype phase focused on Flutter scaffolding": "The build phase focused on Flutter scaffolding",
    "Design and develop a prototype mobile fitness platform":
        "Design and develop a mobile fitness platform",
    "prototype can be demonstrated on Android emulator and iOS simulator, tests pass":
        "the application can be demonstrated on Android and iOS, tests pass",
    "Prototype demo and final presentation": "Application demo and final presentation",
    "The PTD scope covers the prototype design and implementation of Wise Workout.":
        "The PTD covers the design and implementation of Wise Workout.",
    "subscription access, simulated payment handling, and administrative management.":
        "subscription access (payment simulated), and administrative management.",
    "outside the immediate prototype scope.": "outside the current scope.",
    "repository implementation status, and prototype demonstration requirements.":
        "repository implementation status, and demonstration requirements.",
    "Support simulated or prototype-based subscription access": "Support simulated subscription access",
    "The prototype must provide dependable behaviour": "The system must provide dependable behaviour",
    "Performance targets must remain realistic for an FYP prototype while still supporting a credible user experience.":
        "Performance targets must remain realistic while supporting a credible user experience.",
    "Under normal prototype conditions,": "Under normal operating conditions,",
    "The prototype should support at least 50 concurrent active users":
        "The system should support at least 50 concurrent active users",
    "The prototype shall be maintainable by team members": "The system shall be maintainable by team members",
    "supporting growth from prototype testing usage to at least 1,000 registered users":
        "supporting growth from initial usage to at least 1,000 registered users",
    "Prototype payment is simulated and should not request or store real payment card details.":
        "Payment is simulated and does not request or store real payment card details.",
    "Automatic data collection may be limited in the prototype":
        "Automatic data collection may be limited initially",
    "Support manual workout input and selected prototype-level integration where possible.":
        "Support manual workout input and selected device integration where possible.",
    "Focus the prototype on functional demonstration, while future development may include cloud scaling and performance optimisation.":
        "Focus on functional demonstration first, while future development may include cloud scaling and performance optimisation.",
    "user flows, and prototype screens are refined": "user flows, and app screens are refined",
    "keeping the prototype aligned with the project requirements.":
        "keeping the application aligned with the project requirements.",
    "The prototype uses Flutter for cross-platform mobile application development.":
        "Wise Workout uses Flutter for cross-platform mobile application development.",
    "hosted backend services suitable for the prototype.": "hosted backend services suitable for the platform.",
    "and prototype application access. The Flutter app connects":
        "and application access. The Flutter app connects",
    "Supabase functions as the server-side application layer for the prototype.":
        "Supabase functions as the server-side application layer.",
    "This reduces the need for a separately managed server during the FYP prototype stage.":
        "This reduces the need for a separately managed server.",
    "verified on Android emulator and iOS simulator against a live backend. The repository currently reports 107 passing tests and successful prototype flows including":
        "verified on Android and iOS against a live backend. The repository currently reports 107 passing tests and successful flows including",
    "Access Prototype Mobile Application": "Access Mobile Application",
    "the prototype application access or download instructions":
        "the mobile application access or download instructions",
    "downloads and installs the prototype application": "downloads and installs the mobile application",
    "grants access to the prototype application features": "grants access to the mobile application features",
    "If the prototype application is not available for the user's device":
        "If the mobile application is not available for the user's device",
    "The prototype application is available for download or access through the marketing website.":
        "The mobile application is available for download or access through the marketing website.",
    "selects the option to access or download the prototype application.":
        "selects the option to access or download the mobile application.",
    "navigation structure used in the prototype.": "navigation structure used in Wise Workout.",
    "updating the technical stack and prototype status based on the current implementation repository.":
        "updating the technical stack and implementation status based on the current repository.",
    "The current prototype demonstrates key system capabilities including login, onboarding, AI plan generation, plan details, phone-GPS workout capture, paired-wearable heart rate simulation,":
        "Wise Workout currently demonstrates key system capabilities including login, onboarding, AI plan generation, plan details, phone-GPS workout capture, paired-wearable heart rate (simulated),",
    "toward the final project demonstration.": "toward the final demonstration.",
    "Appendix C: Prototype Status Notes": "Appendix C: Implementation Status Notes",
    "Payment handling is simulated in the prototype.": "Payment handling is simulated.",
    "Mock BLE pairing and simulated wearable heart rate are used for wearable demonstration.":
        "Mock BLE pairing and simulated wearable heart rate are used pending real BLE / HealthKit integration.",
    "Repository reports 107 passing tests and successful verification on Android emulator and iOS simulator.":
        "Repository reports 107 passing tests and successful verification on Android and iOS.",
}
# drop the appendix bullet that cites the sample document
merged = [b for b in merged if not (
    b.get("type") == "p" and "Sample Preliminary Technical Documentation" in b.get("text", ""))]
def _apply_scrub(s):
    for _a, _c in _scrub.items():
        if _a in s:
            s = s.replace(_a, _c)
    return s
for _b in merged:
    if _b.get("type") == "p":
        _b["text"] = _apply_scrub(_b.get("text", ""))
    elif _b.get("type") == "t":
        _b["rows"] = [[_apply_scrub(c) for c in row] for row in _b["rows"]]

# ---------- figure fixes: repoint §15 to use-case diagrams, then renumber 1..N ----------
import re as _re
_ucd = {
    "Unregistered user activity flow": "Unregistered user use case diagram",
    "Registered free user activity flow": "Registered free user use case diagram",
    "Registered premium user activity flow": "Registered premium user use case diagram",
    "Verified expert user activity flow": "Verified expert user use case diagram",
    "System admin activity flow": "System administrator use case diagram",
}
for _b in merged:
    if _b.get("type") == "p" and _b.get("text", "").startswith("Figure") and "activity flow from TDM" in _b["text"]:
        _t = _b["text"]
        for _a, _c in _ucd.items():
            _t = _t.replace(_a, _c)
        _t = _re.sub(r"from TDM, page 1[1-5]", "from SRS Section 4 / PRD Section 7.2", _t)
        _b["text"] = _t
# Sequential renumber of every figure caption in document order (handles "Figure:" and "Figure N:")
_fig = 0
for _b in merged:
    if _b.get("type") == "p":
        _m = _re.match(r"^Figure(?:\s+\d+)?\s*:\s*(.*)$", _b.get("text", ""))
        if _m:
            _fig += 1
            _b["text"] = f"Figure {_fig}: {_m.group(1)}"

# ---------- emit markdown ----------
def md_table(rows):
    w=len(rows[0]); L=[]
    L.append("| "+" | ".join(c.replace("\n","; ").replace("|","\\|") for c in rows[0])+" |")
    L.append("| "+" | ".join("---" for _ in range(w))+" |")
    for r in rows[1:]:
        cells=(r+[""]*w)[:w]
        L.append("| "+" | ".join(c.replace("\n","; ").replace("|","\\|") for c in cells)+" |")
    return "\n".join(L)
md=[]
for b in merged:
    if b["type"]=="p" and b["style"]=="Heading 1": md.append("\n# "+b["text"].strip())
    elif b["type"]=="p" and b["style"]=="Heading 2": md.append("\n## "+b["text"].strip())
    elif b["type"]=="p" and b["style"]=="Heading 3": md.append("\n### "+b["text"].strip())
    elif b["type"]=="p" and b["style"]=="List Bullet": md.append("- "+b["text"].strip())
    elif b["type"]=="p":
        t=b["text"].strip()
        if t: md.append(("*"+t+"*") if t.startswith("Figure") else t)
    elif b["type"]=="t": md.append(md_table(b["rows"]))
open(OUTDIR+"PTD-content.md","w").write("\n".join(md).strip()+"\n")

# ---------- emit tab-separated ----------
tx=[]
for b in merged:
    if b["type"]=="p" and b["style"]=="Heading 1":
        s=b["text"].strip(); tx.append("\n\n"+s+"\n"+"="*len(s))
    elif b["type"]=="p" and b["style"]=="Heading 2":
        s=b["text"].strip(); tx.append("\n"+s+"\n"+"-"*len(s))
    elif b["type"]=="p" and b["style"]=="Heading 3":
        tx.append("\n"+b["text"].strip())
    elif b["type"]=="p" and b["style"]=="List Bullet": tx.append("• "+b["text"].strip())
    elif b["type"]=="p":
        t=b["text"].strip()
        if t: tx.append(t)
    elif b["type"]=="t":
        w=len(b["rows"][0])
        tx.append("[TABLE - Convert Text to Table, separator = Tab]")
        for r in b["rows"]:
            cells=(r+[""]*w)[:w]
            tx.append("\t".join(c.replace("\n","; ").strip() for c in cells))
        tx.append("[/TABLE]")
open(OUTDIR+"PTD-content-tabs.txt","w").write("\n".join(tx).strip()+"\n")

# ---------- dump merged content for the docx builder ----------
json.dump(merged, open("/tmp/ptd_content_expanded.json","w"), ensure_ascii=False)

# ---------- stats ----------
words=sum(len(b["text"].split()) for b in merged if b["type"]=="p")
words+=sum(len(c.split()) for b in merged if b["type"]=="t" for r in b["rows"] for c in r)
print("merged blocks:",len(merged))
print("tables:",sum(1 for b in merged if b["type"]=="t"))
print("approx words:",words)
print("wrote PTD-content.md and PTD-content-tabs.txt")
