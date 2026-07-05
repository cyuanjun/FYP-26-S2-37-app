#!/usr/bin/env python3
"""Rebuild the complete PTD content into the v1 template's format (A4, auto-numbered
headings, PTD header). Content comes from /tmp/ptd_content.json (the finished doc);
the format/template comes from the v1 skeleton."""
import shutil, json, re
from copy import deepcopy
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

V1  = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PTD/FYP-26-S2-37-PTD-v1.docx"
OUT = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PTD/FYP-26-S2-37-PTD-v1-FILLED.docx"
shutil.copyfile(V1, OUT)
doc = docx.Document(OUT)
import os
_src = "/tmp/ptd_content_expanded.json" if os.path.exists("/tmp/ptd_content_expanded.json") else "/tmp/ptd_content.json"
blocks = json.load(open(_src))

# ---- capture reference numPr (auto-numbering) from existing v1 headings ----
ref = {"Heading 1": None, "Heading 2": None}
for p in doc.paragraphs:
    if p.style.name in ref and ref[p.style.name] is None:
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn('w:numPr')) is not None:
            ref[p.style.name] = deepcopy(pPr.find(qn('w:numPr')))
assert ref["Heading 1"] is not None and ref["Heading 2"] is not None, "no ref numPr"

# ---- wipe body from 'Introduction' H1 onward (keep final sectPr) ----
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
intro = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and p.text.strip() == "Introduction":
        intro = p._p; break
assert intro is not None
children = list(body)
for el in children[children.index(intro):]:
    if el.tag == qn('w:sectPr'):
        continue
    body.remove(el)

# ---- builders that place each new element just before sectPr ----
def place(el):
    sectPr.addprevious(el)

def strip_num(t):
    return re.sub(r'^\d+(?:\.\d+)*\.?\s+', '', t).strip()

def add_heading(title, level):
    style = f"Heading {level}"
    if style not in [s.name for s in doc.styles]:
        style = "Heading 2"
    p = doc.add_paragraph(style=style)
    p.add_run(strip_num(title))
    # auto-numbering: clone ref numPr if we have one for this level, UNLESS Appendix sub-heading
    refnp = ref.get(f"Heading {level}")
    if refnp is not None and not title.startswith("Appendix "):
        p._p.get_or_add_pPr().append(deepcopy(refnp))
    place(p._p)

def add_paragraph(text, style="Normal"):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if text.startswith("Figure"):
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    place(p._p)

def _set_cell(cell, text):
    parts = str(text).split("\n")
    cell.text = parts[0]
    for extra in parts[1:]:
        cell.add_paragraph(extra)

def add_table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            _set_cell(c, val)
            if i == 0:
                for pp in c.paragraphs:
                    for rr in pp.runs:
                        rr.bold = True
    place(t._tbl)

# ---- optional Table of Contents (non-numbered title + TOC field) ----
def add_toc():
    title = doc.add_paragraph(style="Heading 1")
    title.add_run("Table of Contents")          # no numPr -> not auto-numbered
    place(title._p)
    p = doc.add_paragraph()
    r = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'TOC \\o "1-2" \\h \\z \\u'
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    tx = OxmlElement('w:t'); tx.text = 'Right-click and choose "Update Field" to build the Table of Contents.'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for x in (b, i, s, tx, e):
        r._r.append(x)
    place(p._p)

add_toc()

# ---- emit all content sections in order ----
cur_blocks = None
for b in blocks:
    if b["type"] == "p" and b["style"].startswith("Heading"):
        lvl = int(b["style"][-1])
        add_heading(b["text"].strip(), lvl)
    elif b["type"] == "p":
        style = "List Bullet" if b["style"] == "List Bullet" else "Normal"
        if b["text"].strip():
            add_paragraph(b["text"], style)
    elif b["type"] == "t":
        add_table(b["rows"])

# ---- fix header / footer / cover wording ----
def fix_text(paragraphs, repls):
    for p in paragraphs:
        for run in p.runs:
            for a, c in repls.items():
                if a in run.text:
                    run.text = run.text.replace(a, c)

repls = {
    "Technical Design Manual  (TDM)": "Preliminary Technical Document (PTD)",
    "Technical Design Manual (TDM)": "Preliminary Technical Document (PTD)",
    "19th May 2026": "13th Jun 2026",
    "21st May 2026": "13th Jun 2026",
    "—": "-",  # em dash -> hyphen (cover/template parts; content handled in expand step)
    "–": "-",  # en dash -> hyphen
}
for sec in doc.sections:
    fix_text(sec.header.paragraphs, repls)
    fix_text(sec.footer.paragraphs, repls)
fix_text(doc.paragraphs, repls)   # cover page

# ---- force Word to update fields (TOC) on open ----
settings = doc.settings.element
if settings.find(qn('w:updateFields')) is None:
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
    settings.append(uf)

doc.save(OUT)
print("Saved:", OUT)

# ---- verify ----
chk = docx.Document(OUT)
h1 = [p.text.strip() for p in chk.paragraphs if p.style.name == "Heading 1" and p.text.strip()]
print("H1 sections:", len(h1))
for h in h1:
    print("  -", h)
print("tables:", len(chk.tables))
print("placeholders:", sum(1 for p in chk.paragraphs if "Subsection" in p.text and "content" in p.text))
