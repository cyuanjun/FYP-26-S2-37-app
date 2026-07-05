#!/usr/bin/env python3
"""Build the Preliminary User Manual (PUM).

Single source of truth for PUM content lives in CONTENT below. The script:
  1. Emits PUM-content.md (markdown, mirrors the PTD's PTD-content.md).
  2. Fills the v1 PUM Word template into FYP-26-S2-37-PUM-v1-FILLED.docx
     (A4, auto-numbered headings, PUM header/footer already in template).

The PUM is screenshot-heavy: every screen sub-section emits a [Screenshot: ...]
placeholder + a numbered Figure caption. The team drops the real app/TDM §7
screenshot in place of each placeholder. Figures consistent with the
reconciliation log (premium = $9.99/mo, payment simulated; AI = summaries +
plan suggestions only).
"""
import shutil, re, os
from copy import deepcopy
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Folder of real app screenshots. A fig embeds <slug(screen)>.png from here if present,
# else it falls back to a "[ Insert screenshot ]" placeholder.
SHOTS = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PUM/screenshots"

def slug(s):
    s = s.lower().replace(" screen", "").replace(" wizard", "")
    return re.sub(r"[^a-z0-9]+", "-", s).strip("-")

def nd(s):
    """Normalise em/en dashes to hyphens (team convention, matches the PTD)."""
    return str(s).replace("—", "-").replace("–", "-")

TEMPLATE = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PUM/FYP-26-S2-37-PUM-v1.docx"
OUT_DOCX = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PUM/FYP-26-S2-37-PUM-v1-FILLED.docx"
OUT_MD   = "/Users/cyj/Documents/UOW/FYP/FYP_docs/Submissions/PUM/PUM-content.md"

# ---------------------------------------------------------------------------
# Content model. Each block is a dict:
#   {"k":"h","lvl":1|2|3,"t":title}     heading (numbered, except Document Version Control)
#   {"k":"p","t":text}                   normal paragraph
#   {"k":"b","items":[...]}              bullet list
#   {"k":"n","items":[...]}              numbered (ordered) list
#   {"k":"tbl","rows":[[...],...]}       table (row 0 = header)
#   {"k":"fig","screen":...,"cap":...}   screenshot placeholder + Figure caption
# ---------------------------------------------------------------------------

CONTENT = [
    # ===================== Document Version Control =====================
    {"k": "h", "lvl": 1, "t": "Document Version Control"},
    {"k": "p", "t": "This Preliminary User Manual is a controlled document. The table below records its revision history."},
    {"k": "tbl", "rows": [
        ["Version", "Date", "Description", "Section(s) Affected", "Changed By"],
        ["1.0", "13 Jun 2026", "Initial preliminary user manual", "All", "All team members"],
    ]},

    # ===================== 1. Introduction =====================
    {"k": "h", "lvl": 1, "t": "Introduction"},
    {"k": "p", "t": "Wise Workout is a cross-platform mobile fitness application for Android and iOS. It helps users record workouts using their phone sensors, track their progress over time, and receive AI-assisted progress summaries and plan suggestions. Users set a fitness goal, follow generated training plans, record their sessions, and review their workout history and analytics over time."},

    {"k": "h", "lvl": 2, "t": "What This Manual Covers"},
    {"k": "p", "t": "This manual introduces the application's main features and walks through the key screens a user encounters, in the order they are typically met: getting started and signing in, setting up a fitness profile and goals, following training plans, recording and reviewing workouts, and managing the user profile and settings. Each screen is shown with a screenshot followed by a short description of what the screen does and how to use it."},

    {"k": "h", "lvl": 2, "t": "Intended Audience"},
    {"k": "p", "t": "This manual is written for end users of Wise Workout — primarily registered users of the mobile app. It assumes the reader is comfortable operating a smartphone (installing an app, granting permissions, and using a touch interface) but has no prior knowledge of the application. No technical or development background is required."},

    {"k": "h", "lvl": 2, "t": "Scope and Purpose"},
    {"k": "p", "t": "The purpose of this Preliminary User Manual is to help a new user understand and operate the core features of Wise Workout. Because this is a preliminary submission, the screens shown reflect the features implemented in the current build; additional features and their screens will be added in later releases. Installation paths are described as planned; final store links will be confirmed at public release. The wearable heart-rate capture shown is simulated in the current build and is noted where relevant, so the reader has an accurate picture of the application's state."},

    # ===================== 2. Installation Instructions =====================
    {"k": "h", "lvl": 1, "t": "Installation Instructions"},
    {"k": "p", "t": "The steps below describe how to install and first launch Wise Workout. As this is a preliminary release, the application is distributed for testing rather than through the public app stores; the final store listings will be confirmed at release."},

    {"k": "h", "lvl": 2, "t": "Prerequisites"},
    {"k": "b", "items": [
        "An Android phone running Android 8.0 (Oreo) or later, or an iPhone running iOS 14 or later.",
        "An active internet connection (Wi-Fi or mobile data) for signing in, AI features, and the social feed.",
        "A Wise Workout account, created on the marketing website. A free account is sufficient to use the core features.",
        "Approximately 100 MB of free storage for the app and its data.",
    ]},

    {"k": "h", "lvl": 2, "t": "Installing on Android"},
    {"k": "n", "items": [
        "Open the Wise Workout marketing website (fyp-26-s2-37-website.vercel.app) and create an account, or open the Download section.",
        "Tap Download for Android to download the application package (.apk, approximately 45 MB).",
        "If prompted, allow installation from your browser or from unknown sources in your device settings.",
        "Open the downloaded file and tap Install.",
        "Launch Wise Workout and sign in with your account.",
    ]},

    {"k": "h", "lvl": 2, "t": "Installing on iOS"},
    {"k": "n", "items": [
        "During the project phase the iOS build is distributed for testing (for example, through TestFlight or a development build) rather than the public App Store.",
        "Accept the test invitation and follow the prompt to install Wise Workout.",
        "Open the app and sign in with your account.",
    ]},

    {"k": "h", "lvl": 2, "t": "First-Run Permissions"},
    {"k": "p", "t": "On first use, the application requests the following permissions:"},
    {"k": "b", "items": [
        "Location — to record GPS distance and route for outdoor workouts.",
        "Motion & Fitness — to count steps and detect activity.",
        "Notifications — to deliver workout reminders and progress summaries.",
    ]},
    {"k": "p", "t": "Granting these permissions enables full automatic tracking. They can be changed later in your device settings. If you prefer not to grant sensor permissions, workouts can still be logged manually."},

    # ===================== 3. Key Features =====================
    {"k": "h", "lvl": 1, "t": "Key Features"},
    {"k": "p", "t": "The table below summarises what a user can do with Wise Workout in the current build."},
    {"k": "tbl", "rows": [
        ["Feature", "What the user can do"],
        ["Workout tracking", "Record workouts using phone sensors (GPS plus step and motion) or enter them manually; start a suggested session from a plan or a freeform session on the spot."],
        ["Progress & history", "View workout history and basic analytics by day, week and month, including a comparison with the previous week."],
        ["AI-assisted support", "Receive AI progress summaries and AI plan suggestions. Plans follow the user's selected goal timeline; premium plans are personalised from the user's fitness profile."],
        ["Fitness profile & goals", "Set body metrics, training experience, preferences, and a primary goal with a target and timeline; track goal progress over time."],
        ["Connected devices", "Pair a wearable heart-rate monitor (simulated in the current build) so its readings are recorded into a session."],
        ["Notification preferences", "Choose which reminders and notifications to receive from the notification settings screen."],
    ]},

    # ===================== 4. Initial GUIs — Screen Walkthrough =====================
    {"k": "h", "lvl": 1, "t": "Initial GUIs — Screen Walkthrough"},
    {"k": "p", "t": "This section walks through the application screen by screen. Each sub-section shows a screenshot followed by a short description of the screen and how to use it. The screens are grouped to follow the journey of a typical user, from first launch through everyday use. This preliminary manual covers the screens implemented in the current build; remaining screens (the expert marketplace, the social feed, the premium upgrade flow, and the expert and administrator portals) will be documented as those features are completed."},

    # ---- 4.1 Getting Started ----
    {"k": "h", "lvl": 2, "t": "Getting Started"},
    {"k": "h", "lvl": 3, "t": "Splash Screen"},
    {"k": "fig", "screen": "Splash screen", "cap": "Splash screen shown while the app loads."},
    {"k": "p", "t": "When you open the app, the WISE WORKOUT wordmark and the tagline \"Train smart. Move better.\" appear briefly while the app starts up and checks whether you are already signed in."},
    {"k": "howto", "items": [
        "No action is needed — this screen appears automatically when the app launches.",
        "After about a second it moves on by itself: to Home if you are already signed in, to the onboarding wizard if you are a new user, or to the Sign In screen otherwise.",
    ]},

    {"k": "h", "lvl": 3, "t": "Sign In"},
    {"k": "fig", "screen": "Login screen", "cap": "Sign-in screen."},
    {"k": "p", "t": "The sign-in screen shows the WISE WORKOUT wordmark, an EMAIL field, a PASSWORD field, a full-width LOG IN button, a \"Forgot password?\" link, and a line at the bottom inviting new users to sign up on the Wise Workout website."},
    {"k": "howto", "items": [
        "Enter your registered email address in the EMAIL field.",
        "Enter your password in the PASSWORD field.",
        "Tap LOG IN to sign in — you are taken to Home, or to the onboarding wizard on a first login.",
        "Tap \"Forgot password?\" to reset your password, or the sign-up line to register on the website.",
    ]},

    {"k": "h", "lvl": 3, "t": "Forgot Password"},
    {"k": "fig", "screen": "Forgot password screen", "cap": "Password recovery."},
    {"k": "p", "t": "This screen has a \"FORGOT PASSWORD\" title, a short instruction line, a single email field, a \"SEND RESET LINK\" button, and a \"Back to log in\" link."},
    {"k": "howto", "items": [
        "Enter the email address linked to your account.",
        "Tap SEND RESET LINK — a secure reset link (valid for 30 minutes) is emailed to you.",
        "Open the email and tap the link to choose a new password.",
        "Tap \"Back to log in\" to return to the Sign In screen.",
    ]},

    {"k": "h", "lvl": 3, "t": "Onboarding"},
    {"k": "fig", "screen": "Onboarding wizard", "cap": "First-login onboarding wizard."},
    {"k": "p", "t": "On your first sign-in a five-step wizard sets up your profile, with a progress bar across the top and Back/Continue buttons at the bottom. The steps are: a welcome screen; About You (date of birth, sex, height, weight); How You Train (activity level, training experience, and preferred workout types); Your Goal (a goal card, target, days per week, and timeline); and a final step that generates your first plan."},
    {"k": "howto", "items": [
        "On the welcome step, tap LET'S GO.",
        "On About You, enter your date of birth, sex, height and weight.",
        "On How You Train, pick your activity level and experience, and tap the workout types you prefer (or \"+ Add your own\").",
        "On Your Goal, choose a goal card and set the target, days per week, and timeline.",
        "Tap GENERATE MY PLAN, wait for your plan to be created, then tap START TRAINING.",
        "Use Back at any point to revise an earlier step.",
    ]},

    # ---- 4.2 Home Dashboard ----
    {"k": "h", "lvl": 2, "t": "Home Dashboard"},
    {"k": "fig", "screen": "Home dashboard", "cap": "Home dashboard."},
    {"k": "p", "t": "Home greets you by name, shows whether you are a Free or Premium member, and presents a \"Get moving\" card that points you to Train to record a workout and to History for your stats and AI summary. A five-tab bar runs along the bottom — Home, Experts, Train, Social and History — and your avatar in the top-right opens your Profile."},
    {"k": "howto", "items": [
        "Read the greeting and your member tier at the top of the screen.",
        "Tap Train in the bottom bar to record a workout, or History to review past sessions.",
        "Tap your avatar (top-right) to open Profile.",
        "Use the bottom tabs to move between Home, Experts, Train, Social and History.",
    ]},

    # ---- 4.3 Training & Plans ----
    {"k": "h", "lvl": 2, "t": "Training & Plans"},
    {"k": "h", "lvl": 3, "t": "Train"},
    {"k": "fig", "screen": "Train screen", "cap": "Train screen."},
    {"k": "p", "t": "Train is your training home. Under AI SUGGESTED PLAN, a card shows your active plan's name and summary (for example \"12-Week Lose Weight Plan · 4x per week\"), today's workout highlighted with its duration, and chips for the rest of the week's workouts. A \"VIEW PLANS\" link sits at the top-right. Below, a DEVICES section shows your connected wearable with an \"+ ADD DEVICE\" button, and a \"START FREEFORM WORKOUT\" button is pinned at the bottom."},
    {"k": "howto", "items": [
        "Tap today's highlighted workout in the plan card to open and start it.",
        "Tap \"VIEW PLANS\" (top-right) to open My Plans.",
        "Tap your device, or \"+ ADD DEVICE\", to manage connected devices.",
        "Tap START FREEFORM WORKOUT to record an unplanned session at any time.",
    ]},

    {"k": "h", "lvl": 3, "t": "My Plans"},
    {"k": "fig", "screen": "My Plans screen", "cap": "My Plans."},
    {"k": "p", "t": "My Plans lists every plan you have generated, grouped into ACTIVE (your current plan, marked with an \"ACTIVE\" badge) and SAVED (older plans, newest first). Each card shows the plan name, its length and weekly frequency, and a \"PERSONALISED\" badge for premium plans."},
    {"k": "howto", "items": [
        "Tap any plan card to open its full schedule in Plan Detail.",
        "If you have no plans yet, tap \"Set a goal\" to create your first AI plan.",
    ]},

    {"k": "h", "lvl": 3, "t": "Plan Detail"},
    {"k": "fig", "screen": "Plan Detail screen", "cap": "Plan detail with week-by-week schedule."},
    {"k": "p", "t": "Plan Detail shows the plan's name, length, frequency and AI description, then the schedule laid out week by week. Each week lists its workouts as rows showing the day, workout name and duration; today's row is highlighted. A \"Start today's workout\" button and a \"Regenerate plan\" link sit at the bottom. Generated plans follow the goal timeline you set, so a twelve-week goal produces a twelve-week plan."},
    {"k": "howto", "items": [
        "Scroll through the weeks to see the full schedule.",
        "Tap any workout row to open its details (type, duration and description).",
        "When viewing today's workout, tap Start to begin the session.",
        "If this plan is not your active one, tap \"Use This Plan\" to make it active.",
        "Tap \"Regenerate plan\" to replace it with a fresh AI-generated plan.",
    ]},

    {"k": "h", "lvl": 3, "t": "Connected Devices"},
    {"k": "fig", "screen": "Connected Devices screen", "cap": "Connected devices and wearable pairing."},
    {"k": "p", "t": "This screen lists the wearables paired to your account, each row showing the device name, its connection status and when it last synced, plus an \"+ Add device\" button at the bottom. In the current build, wearable pairing and heart-rate data are simulated for demonstration."},
    {"k": "howto", "items": [
        "Tap \"+ Add device\" to scan for and pair a wearable heart-rate monitor.",
        "Tap a paired device to manage it.",
        "Once paired, the device's heart rate is recorded into your workouts; sessions logged without a device use phone sensors or manual entry.",
    ]},

    # ---- 4.4 Recording a Workout ----
    {"k": "h", "lvl": 2, "t": "Recording a Workout"},
    {"k": "h", "lvl": 3, "t": "Active Workout"},
    {"k": "fig", "screen": "Active workout screen", "cap": "Live workout capture."},
    {"k": "p", "t": "The live session screen shows a large TIME timer in the centre with metric tiles below it — for a cardio activity, Distance and Pace (and heart rate when a device is connected). An activity label such as \"Running\" sits at the bottom-left, with a large circular START button in the centre and a \"TAP TO BEGIN\" prompt. Once you begin, the timer and metrics go live, the button becomes PAUSE, and a red END button appears."},
    {"k": "howto", "items": [
        "Tap the large central button to START — the timer and metrics go live.",
        "Tap PAUSE to pause and RESUME to continue.",
        "Tap the activity label (bottom-left) to change the workout type.",
        "Tap the red END button when you finish, then confirm \"Save & Finish\" to see your summary.",
    ]},

    {"k": "h", "lvl": 3, "t": "Workout Complete & Share"},
    {"k": "fig", "screen": "Workout summary screen", "cap": "Workout-complete summary."},
    {"k": "p", "t": "After ending a session, the Workout Complete screen shows a green checkmark, the XP you earned and your current streak, and a stats grid (duration, distance, pace and workout type; calories and heart-rate stats appear when available). Below are a \"name this workout\" field, \"How did it feel?\" mood chips (Great / Good / Okay / Tough), a \"Share to Social\" toggle, a private notes box, and \"Save & Finish\" / \"Skip\" buttons at the bottom."},
    {"k": "howto", "items": [
        "Optionally type a name for the workout in the name field.",
        "Tap a \"How did it feel?\" chip (Great / Good / Okay / Tough).",
        "Turn on \"SHARE TO SOCIAL\" and add a caption to post it — you can also share to Facebook, Instagram, Twitter or TikTok.",
        "Add any private notes (these are never shared with anyone).",
        "Tap \"Save & Finish\" to save, or \"Skip\" to keep the session without the extra details.",
    ]},

    # ---- 4.5 History ----
    {"k": "h", "lvl": 2, "t": "History"},
    {"k": "h", "lvl": 3, "t": "Workout History"},
    {"k": "fig", "screen": "History screen", "cap": "Workout history and basic analytics."},
    {"k": "p", "t": "History shows a Basic Workout Analytics card with Day / Week / Month period pills and tiles for sessions, active minutes, calories and average/max heart rate, each with an arrow comparing it to the previous period. Below it your past sessions are listed and grouped by week (This Week / Last Week / Earlier), each card showing the workout name, date, duration and key metrics. Free accounts see the current month of history."},
    {"k": "howto", "items": [
        "Tap a period pill (Day / Week / Month) to change the time window for the analytics and the list.",
        "Scroll the session list, grouped by week with the newest first.",
        "Tap any session card to open its full detail.",
    ]},

    {"k": "h", "lvl": 3, "t": "Workout Detail"},
    {"k": "fig", "screen": "History detail screen", "cap": "Individual workout detail."},
    {"k": "p", "t": "This read-only recap of a past workout shows its name, date and duration, a stats grid (duration, distance, pace, calories, average and max heart rate), how the session felt, and any private notes. For sessions recorded with live heart-rate or GPS data, graphs you can switch between (heart rate, pace and more) also appear. An \"Edit\" button in the top-right lets you change the name, mood, notes or sharing."},
    {"k": "howto", "items": [
        "Review the stats, how the session felt, and your notes.",
        "If the session was recorded with heart-rate or GPS data, tap the graph pills to switch metrics.",
        "Tap \"Edit\" to change the name, mood, notes or sharing, then save your changes.",
        "Tap Back to return to the History list.",
    ]},

    # ---- 4.6 Profile, Goals & Settings ----
    {"k": "h", "lvl": 2, "t": "Profile, Goals & Settings"},
    {"k": "h", "lvl": 3, "t": "Profile"},
    {"k": "fig", "screen": "Profile screen", "cap": "Profile hub."},
    {"k": "p", "t": "Profile shows your avatar, name, handle and level/XP at the top, a row of stats (workouts, active days, weekly streak), and a menu: Account Settings, Fitness Profile, Fitness Goals, Notifications and Submit Feedback. A \"GO PREMIUM\" pill appears in the top-right for free users, and a \"Log out\" button sits at the bottom."},
    {"k": "howto", "items": [
        "Tap a menu row to open Account Settings, Fitness Profile, Fitness Goals, Notifications or Submit Feedback.",
        "Tap the pencil on your avatar to change your photo.",
        "Tap \"GO PREMIUM\" to upgrade (free users only).",
        "Tap \"Log out\" to sign out of the app.",
    ]},

    {"k": "h", "lvl": 3, "t": "Fitness Profile"},
    {"k": "fig", "screen": "Fitness profile screen", "cap": "Fitness profile."},
    {"k": "p", "t": "Fitness Profile holds the details that personalise your plans and analytics: BODY METRICS (date of birth, sex, height, weight), ACTIVITY LEVEL, TRAINING EXPERIENCE, and chip lists for PREFERRED WORKOUTS, DIET, ALLERGIES and INJURIES / LIMITATIONS. A \"SAVE PROFILE\" button is at the bottom."},
    {"k": "howto", "items": [
        "Tap a body-metric row (for example Height) to edit its value.",
        "Tap the Activity Level card and pick a level; tap a Training Experience chip to select it.",
        "Tap the \"+\" next to Preferred Workouts, Diet, Allergies or Injuries to add items, or tap a chip to toggle it.",
        "Tap \"SAVE PROFILE\" to save your changes.",
    ]},

    {"k": "h", "lvl": 3, "t": "Fitness Goals"},
    {"k": "fig", "screen": "Fitness goals screen", "cap": "Fitness goals."},
    {"k": "p", "t": "Fitness Goals lets you set your PRIMARY GOAL by tapping a card — Lose Weight, Build Muscle, Improve Endurance or Maintain Fitness, each with a short descriptor and a checkmark on the selected one — then set your WEEKLY COMMITMENT with a days-per-week stepper. A \"SAVE GOAL\" button is at the bottom."},
    {"k": "howto", "items": [
        "Tap a goal card to choose your primary goal.",
        "Use the minus/plus stepper to set how many days per week you will train.",
        "Tap \"SAVE GOAL\" — this saves the goal and generates a matching AI plan.",
    ]},

    {"k": "h", "lvl": 3, "t": "Account Settings"},
    {"k": "fig", "screen": "Account settings screen", "cap": "Account settings."},
    {"k": "p", "t": "Account Settings has a PERSONAL INFO section (full name, username, email), a PREFERENCES section with a METRIC / IMPERIAL unit toggle, and a SECURITY section with a \"CHANGE PASSWORD\" button."},
    {"k": "howto", "items": [
        "Tap a field (for example Full Name or Email) to edit it.",
        "Tap METRIC or IMPERIAL to switch your measurement units — this applies immediately.",
        "Tap \"CHANGE PASSWORD\" to start a password reset.",
    ]},

    {"k": "h", "lvl": 3, "t": "Notification Settings"},
    {"k": "fig", "screen": "Notification settings screen", "cap": "Notification preferences."},
    {"k": "p", "t": "This screen groups notifications into Workout Reminders (daily, missed-workout and inactivity), Summaries (weekly summary), Social (friend activity, likes and comments, challenge invites) and Marketing & Updates, each row with an on/off toggle."},
    {"k": "howto", "items": [
        "Flip any toggle on or off to choose which notifications you receive.",
        "Changes are saved automatically — there is no separate Save button.",
    ]},

    {"k": "h", "lvl": 3, "t": "Submit Feedback"},
    {"k": "fig", "screen": "Submit feedback screen", "cap": "Submit feedback."},
    {"k": "p", "t": "Submit Feedback lets you tell the team about a bug, request a feature, or send general comments. You pick a category tile (Bug / Feature request / General), type your message in the text box (at least 10 characters), and tap the \"Submit feedback\" button; a thank-you screen then confirms it was sent."},
    {"k": "howto", "items": [
        "Tap a category: Bug, Feature request, or General.",
        "Type your message in the box — the Submit button enables once it is long enough.",
        "Tap \"Submit feedback\" to send it.",
        "On the confirmation screen, tap \"Submit another\" or \"Back to Profile\".",
    ]},
]

# ---------------------------------------------------------------------------
# 1) Emit PUM-content.md
# ---------------------------------------------------------------------------
def emit_md():
    fig = 0
    lines = [
        "# PUM — Assembled Content (Preliminary User Manual)",
        "",
        "Single source for the **Wise Workout Preliminary User Manual** body. Generated by "
        "`scripts/build_pum.py`, which also fills the Word template "
        "(`FYP-26-S2-37-PUM-v1-FILLED.docx`). Each `[Screenshot: …]` line marks where the "
        "team drops the matching app or TDM §7 screenshot. Figures are auto-numbered.",
        "",
        "---",
        "",
    ]
    for b in CONTENT:
        k = b["k"]
        if k == "h":
            lines.append("#" * (b["lvl"] + 1) + " " + b["t"])
            lines.append("")
        elif k == "p":
            lines.append(b["t"]); lines.append("")
        elif k == "b":
            for it in b["items"]:
                lines.append(f"- {it}")
            lines.append("")
        elif k == "n":
            for i, it in enumerate(b["items"], 1):
                lines.append(f"{i}. {it}")
            lines.append("")
        elif k == "howto":
            lines.append("**How to use:**")
            lines.append("")
            for it in b["items"]:
                lines.append(f"- {it}")
            lines.append("")
        elif k == "tbl":
            rows = b["rows"]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for r in rows[1:]:
                lines.append("| " + " | ".join(r) + " |")
            lines.append("")
        elif k == "fig":
            fig += 1
            lines.append(f"> **[Screenshot: {b['screen']}]**")
            lines.append("")
            lines.append(f"*Figure {fig}: {b['cap']}*")
            lines.append("")
    with open(OUT_MD, "w") as f:
        f.write("\n".join(lines))
    print("Saved:", OUT_MD, f"({fig} figures)")
    return fig

# ---------------------------------------------------------------------------
# 2) Fill the Word template
# ---------------------------------------------------------------------------
def emit_docx():
    shutil.copyfile(TEMPLATE, OUT_DOCX)
    doc = docx.Document(OUT_DOCX)

    # Capture reference numPr (auto-numbering) from the template's numbered H1/H2.
    ref = {"Heading 1": None, "Heading 2": None}
    for p in doc.paragraphs:
        s = p.style.name
        if s in ref and ref[s] is None:
            pPr = p._p.pPr
            if pPr is not None and pPr.find(qn('w:numPr')) is not None:
                ref[s] = deepcopy(pPr.find(qn('w:numPr')))
    assert ref["Heading 1"] is not None and ref["Heading 2"] is not None, "no ref numPr in template"
    # Synthesize H3 numPr from H2's list (same numId, list level 2) so 4.x.y numbers.
    ref["Heading 3"] = deepcopy(ref["Heading 2"])
    ilvl = ref["Heading 3"].find(qn('w:ilvl'))
    if ilvl is not None:
        ilvl.set(qn('w:val'), "2")

    # Wipe body from the "Introduction" H1 onward (KEEP the cover, Document Version Control,
    # revision table, TOC content-control, and the section break that carries the body header).
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    intro = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "Introduction":
            intro = p._p
            break
    assert intro is not None, "no Introduction H1 in template"
    children = list(body)
    for el in children[children.index(intro):]:
        if el.tag == qn('w:sectPr'):
            continue
        body.remove(el)

    def place(el):
        sectPr.addprevious(el)

    def add_heading(title, level, numbered=True):
        p = doc.add_paragraph(style=f"Heading {level}")
        p.add_run(nd(title))
        refnp = ref.get(f"Heading {level}")
        if numbered and refnp is not None:
            p._p.get_or_add_pPr().append(deepcopy(refnp))
        place(p._p)

    def add_paragraph(text, style="Normal", italic=False, center=False, bold=False):
        p = doc.add_paragraph(style=style)
        run = p.add_run(nd(text))
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        place(p._p)

    def add_table(rows):
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                c = t.rows[i].cells[j]
                c.text = nd(val)
                if i == 0:
                    for pp in c.paragraphs:
                        for rr in pp.runs:
                            rr.bold = True
        place(t._tbl)

    # The template already supplies Document Version Control + revision table + a TOC content
    # control, so skip our CONTENT front matter and start placing at the "Introduction" heading.
    fig = 0
    started = False
    for b in CONTENT:
        k = b["k"]
        if not started:
            if k == "h" and b["t"] == "Introduction":
                started = True
            else:
                continue
        if k == "h":
            add_heading(b["t"], b["lvl"], numbered=True)
        elif k == "p":
            add_paragraph(b["t"])
        elif k == "b":
            for it in b["items"]:
                add_paragraph(it, style="List Bullet")
        elif k == "n":
            for it in b["items"]:
                add_paragraph(it, style="List Number")
        elif k == "howto":
            add_paragraph("How to use:", bold=True)
            for it in b["items"]:
                add_paragraph(it, style="List Bullet")
        elif k == "tbl":
            add_table(b["rows"])
        elif k == "fig":
            fig += 1
            shot = os.path.join(SHOTS, slug(b["screen"]) + ".png")
            if os.path.exists(shot):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(shot, height=Inches(4.2))
                place(p._p)
            else:
                add_paragraph(f"[ Insert screenshot: {b['screen']} ]", center=True)
            add_paragraph(f"Figure {fig}: {b['cap']}", italic=True, center=True)

    # normalise dashes in cover / header / footer (content already clean)
    def fix(paras):
        for p in paras:
            for run in p.runs:
                run.text = run.text.replace("—", "-").replace("–", "-")
    for sec in doc.sections:
        fix(sec.header.paragraphs)
        fix(sec.footer.paragraphs)

    # update fields (TOC) on open
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
        settings.append(uf)

    doc.save(OUT_DOCX)
    print("Saved:", OUT_DOCX, f"({fig} figures)")

    # verify
    chk = docx.Document(OUT_DOCX)
    h1 = [p.text.strip() for p in chk.paragraphs if p.style.name == "Heading 1" and p.text.strip()]
    h2 = [p.text.strip() for p in chk.paragraphs if p.style.name == "Heading 2" and p.text.strip()]
    print("H1:", len(h1), "| H2:", len(h2), "| tables:", len(chk.tables))
    for h in h1:
        print("  -", h)


if __name__ == "__main__":
    emit_md()
    emit_docx()
